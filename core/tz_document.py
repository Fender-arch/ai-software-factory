"""Compose and export the full TZ document from the knowledge graph."""

from __future__ import annotations

import io
import re
import uuid
from pathlib import Path
from typing import Literal

from sqlalchemy.orm import Session

from core.config import get_settings
from core.models import Project
from discovery.tz_outline import plan_from_state, resolve_active_topics
from knowledge.repository import KnowledgeRepository
from knowledge.types import normalize_requirement_status

TzExportFormat = Literal["md", "pdf", "docx"]

STATUS_RU = {
    "new": "новое",
    "processed": "отработано",
    "needs_clarification": "уточняется",
    "conflict": "конфликт",
    "rejected": "отклонено",
    "superseded": "заменено",
}

PRODUCT_RU = {
    "website": "сайт",
    "telegram_bot": "Telegram-бот",
    "rest_service": "REST-сервис",
    "ai_automation": "AI-автоматизация",
    "mobile_native": "нативное приложение",
}

PROJECT_STATUS_RU = {
    "NEW": "новый",
    "INTERVIEW": "интервью",
    "ANALYZING": "анализ",
    "WAITING_CUSTOMER": "ждём заказчика",
    "WAITING_OWNER": "ждём владельца",
    "WAITING_CLIENT_ESTIMATE": "ждём смету клиента",
    "READY": "готов",
    "ARCHIVED": "в архиве",
}

CLIENT_TZ_TITLE = "Техническое задание"

_EXTRA_TZ_SECTIONS_RU = (
    ("closing_additions", "Дополнения заказчика"),
    ("source_brief", "Исходная постановка (файл / нейросеть)"),
    ("owner_review_supplement", "Дополнения после ревью"),
)

_MD_LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")
_MD_MARKUP_RE = re.compile(r"[*`]+")

_FONT_CANDIDATES = (
    Path(__file__).resolve().parent / "fonts" / "DejaVuSans.ttf",
    Path(r"C:\Windows\Fonts\arial.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
)


class TzExportError(ValueError):
    """Could not build a TZ file."""


def _project_state(kg: KnowledgeRepository, project: Project) -> dict:
    entities = kg.list_entities(project.id, type_="Project")
    if not entities:
        return {}
    return dict(entities[0].payload or {})


def heading_anchor(title: str) -> str:
    """GitHub-style slug so Markdown TOC links resolve in common MD viewers."""
    text = (title or "").strip().lower()
    text = re.sub(r"[^\w\s\-а-яё]", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+", "-", text)
    return text.strip("-") or "section"


def plain_markdown_line(line: str) -> str:
    """Drop link markup and emphasis for PDF/DOCX (exporters are line-oriented)."""
    text = _MD_LINK_RE.sub(r"\1", line)
    return _MD_MARKUP_RE.sub("", text)


def resolve_owner_contacts(state: dict) -> dict[str, str]:
    """Studio/owner contacts: Project payload.owner_contacts overrides env settings.

    Hook (no new table): KG ``Project`` entity ``payload.owner_contacts``
    ``{studio, name, email, phone, telegram, note}``. Env fallback:
    ``STUDIO_NAME``, ``OWNER_CONTACT_NAME``, ``OWNER_CONTACT_EMAIL``,
    ``OWNER_CONTACT_PHONE``, ``OWNER_CONTACT_TELEGRAM``.
    ``OWNER_TELEGRAM_ID`` stays HITL-only and is not printed.
    """
    hook = state.get("owner_contacts")
    if not isinstance(hook, dict):
        hook = {}
    settings = get_settings()
    return {
        "studio": str(
            hook.get("studio") or hook.get("studio_name") or settings.studio_name or ""
        ).strip(),
        "name": str(
            hook.get("name") or hook.get("owner_name") or settings.owner_contact_name or ""
        ).strip(),
        "email": str(hook.get("email") or settings.owner_contact_email or "").strip(),
        "phone": str(hook.get("phone") or settings.owner_contact_phone or "").strip(),
        "telegram": str(
            hook.get("telegram") or settings.owner_contact_telegram or ""
        ).strip(),
        "note": str(hook.get("note") or "").strip(),
    }


def _contact_lines_from_entities(entities: list) -> list[str]:
    lines: list[str] = []
    for ent in entities:
        payload = getattr(ent, "payload", None) or {}
        desc = str(payload.get("description") or getattr(ent, "name", "") or "").strip()
        if desc:
            lines.append(desc)
    return lines


def _format_owner_contact_lines(contacts: dict[str, str]) -> list[str]:
    labels = (
        ("studio", "Студия"),
        ("name", "Имя"),
        ("email", "Email"),
        ("phone", "Телефон"),
        ("telegram", "Telegram"),
        ("note", "Комментарий"),
    )
    lines = [f"- {label}: {contacts[key]}" for key, label in labels if contacts.get(key)]
    if not lines:
        return ["- _Не указаны._"]
    return lines


def _req_code(section_no: int, item_no: int) -> str:
    return f"ТЗ-{section_no}.{item_no}"


def _requirement_line(section_no: int, item_no: int, ent) -> str:
    payload = ent.payload or {}
    desc = payload.get("description") or ent.name
    status = STATUS_RU.get(normalize_requirement_status(ent.status), ent.status or "")
    priority = payload.get("priority") or "should"
    code = _req_code(section_no, item_no)
    return f"- **{code}** [{status}] ({priority}) {desc}"


def _plain_item_line(section_no: int, item_no: int, text: str) -> str:
    return f"- **{_req_code(section_no, item_no)}** {text}"


def compose_tz_markdown(db: Session, project: Project) -> str:
    """Client TZ in Russian, derived from current KG (presentation only)."""
    kg = KnowledgeRepository(db)
    state = _project_state(kg, project)
    requirements = [
        e
        for e in kg.list_entities(project.id, type_="Requirement")
        if e.status != "archived"
    ]
    open_questions = kg.list_entities(project.id, type_="OpenQuestion")
    risks = kg.list_entities(project.id, type_="Risk")
    task_shape = state.get("task_shape")
    plan = plan_from_state(state)
    outline = resolve_active_topics(project.product_type, task_shape=task_shape, plan=plan)

    req_by_topic: dict[str, list] = {}
    unscoped: list = []
    for ent in requirements:
        if normalize_requirement_status(ent.status) == "superseded":
            continue
        tid = str((ent.payload or {}).get("topic_id") or "")
        if tid:
            req_by_topic.setdefault(tid, []).append(ent)
        else:
            unscoped.append(ent)

    product = PRODUCT_RU.get(project.product_type or "", project.product_type or "не задан")
    status_key = project.status.value
    status_ru = PROJECT_STATUS_RU.get(status_key, status_key)
    owner = resolve_owner_contacts(state)
    customer_lines = _contact_lines_from_entities(req_by_topic.get("contacts") or [])
    preferred = _contact_lines_from_entities(req_by_topic.get("preferred_contact") or [])
    if project.customer_telegram_id:
        customer_lines.insert(0, f"Telegram заказчика: `{project.customer_telegram_id}`")
    for line in preferred:
        customer_lines.append(f"Предпочтительный канал: {line}")
    if not customer_lines:
        customer_lines = ["_Не указаны._"]

    lines = [
        f"# {CLIENT_TZ_TITLE} — {project.name}",
        "",
        "## Мета",
        "",
        f"- **Проект:** {project.name}",
        f"- **Идентификатор:** `{project.id}`",
        f"- **Тип продукта:** {product} (`{project.product_type or 'unspecified'}`)",
        f"- **Форма задачи:** `{task_shape or '—'}`",
        f"- **Статус проекта:** {status_ru} (`{status_key}`)",
        "",
        "### Контакты заказчика",
        "",
        *[
            item if item.startswith("- ") or item.startswith("_") else f"- {item}"
            for item in customer_lines
        ],
        "",
        "### Контакты исполнителя",
        "",
        *_format_owner_contact_lines(owner),
        "",
    ]

    sections: list[tuple[int, str, list[str]]] = []
    extra_ids = {tid for tid, _ in _EXTRA_TZ_SECTIONS_RU}
    outline_ids = {topic.id for topic in outline}

    for topic in outline:
        section_no = len(sections) + 1
        ents = req_by_topic.get(topic.id) or []
        body: list[str] = []
        if not ents:
            body.append("_Пока не зафиксировано._")
        else:
            for idx, ent in enumerate(ents, start=1):
                body.append(_requirement_line(section_no, idx, ent))
        sections.append((section_no, topic.title_ru, body))

    for extra_id, extra_title in _EXTRA_TZ_SECTIONS_RU:
        ents = req_by_topic.get(extra_id) or []
        if not ents:
            continue
        section_no = len(sections) + 1
        body = [_requirement_line(section_no, idx, ent) for idx, ent in enumerate(ents, start=1)]
        sections.append((section_no, extra_title, body))

    leftover_unscoped = [
        e
        for e in unscoped
        if str((e.payload or {}).get("topic_id") or "") not in extra_ids
        and str((e.payload or {}).get("topic_id") or "") not in outline_ids
    ]
    if leftover_unscoped:
        section_no = len(sections) + 1
        body = [
            _requirement_line(section_no, idx, ent)
            for idx, ent in enumerate(leftover_unscoped, start=1)
        ]
        sections.append((section_no, "Прочие требования", body))

    open_active = [e for e in open_questions if e.status == "open"]
    section_no = len(sections) + 1
    if not open_active:
        q_body = ["_Нет._"]
    else:
        q_body = [
            _plain_item_line(
                section_no,
                idx,
                str((ent.payload or {}).get("question") or ent.name),
            )
            for idx, ent in enumerate(open_active, start=1)
        ]
    sections.append((section_no, "Открытые вопросы", q_body))

    section_no = len(sections) + 1
    if not risks:
        risk_body = ["_Не записаны._"]
    else:
        risk_body = [
            _plain_item_line(
                section_no,
                idx,
                str((ent.payload or {}).get("description") or ent.name),
            )
            for idx, ent in enumerate(risks, start=1)
        ]
    sections.append((section_no, "Риски", risk_body))

    lines.extend(["## Оглавление", ""])
    for number, title, _body in sections:
        heading = f"{number}. {title}"
        lines.append(f"{number}. [{title}](#{heading_anchor(heading)})")
    lines.append("")

    for number, title, body in sections:
        heading = f"{number}. {title}"
        lines.extend([f"## {heading}", ""])
        lines.extend(body)
        lines.append("")

    return "\n".join(lines)


def safe_export_stem(name: str) -> str:
    slug = re.sub(r"[^\w\-]+", "-", (name or "tz").strip(), flags=re.UNICODE)
    return (slug.strip("-") or "tz")[:72]


def _safe_stem(name: str) -> str:
    return safe_export_stem(name)


def _cyrillic_font() -> Path:
    for path in _FONT_CANDIDATES:
        if path.is_file():
            return path
    raise TzExportError("no Unicode TTF found for PDF export")


def markdown_to_docx(markdown: str) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for raw in markdown.splitlines():
        line = plain_markdown_line(raw.rstrip())
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line == "---":
            doc.add_paragraph("—" * 12)
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_to_pdf(markdown: str) -> bytes:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    font = _cyrillic_font()
    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.add_page()
    pdf.add_font("TzSans", fname=str(font))
    pdf.set_text_color(28, 32, 38)
    usable = pdf.epw

    def write(text: str, size: int, height: float) -> None:
        pdf.set_font("TzSans", size=size)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(
            usable,
            height,
            text,
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )

    for raw in markdown.splitlines():
        line = plain_markdown_line(raw.replace("\t", "  "))
        if not line.strip():
            pdf.ln(3)
            continue
        if line.startswith("# "):
            title = line[2:].strip()
            if hasattr(pdf, "start_section"):
                pdf.start_section(title, level=0)
            write(title, 18, 9)
            pdf.ln(2)
        elif line.startswith("## "):
            title = line[3:].strip()
            if hasattr(pdf, "start_section"):
                pdf.start_section(title, level=1)
            write(title, 14, 8)
            pdf.ln(1)
        elif line.startswith("### "):
            title = line[4:].strip()
            if hasattr(pdf, "start_section"):
                pdf.start_section(title, level=2)
            write(title, 12, 7)
        elif line.startswith("- "):
            write(f"- {line[2:].strip()}", 11, 6)
        else:
            write(line, 11, 6)
    return bytes(pdf.output())


def export_markdown_file(
    markdown: str, stem: str, fmt: TzExportFormat
) -> tuple[bytes, str, str]:
    """Turn Markdown into md / pdf / docx bytes (TZ, смета, other artifacts)."""
    safe = safe_export_stem(stem)
    if fmt == "md":
        return markdown.encode("utf-8"), "text/markdown; charset=utf-8", f"{safe}.md"
    if fmt == "docx":
        return (
            markdown_to_docx(markdown),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{safe}.docx",
        )
    if fmt == "pdf":
        return markdown_to_pdf(markdown), "application/pdf", f"{safe}.pdf"
    raise TzExportError(f"unsupported format: {fmt}")


def export_tz_file(
    db: Session, project: Project, fmt: TzExportFormat
) -> tuple[bytes, str, str]:
    markdown = compose_tz_markdown(db, project)
    return export_markdown_file(markdown, project.name, fmt)


def export_client_estimate_file(
    db: Session, project: Project, fmt: TzExportFormat
) -> tuple[bytes, str, str]:
    from core.client_estimate import (
        client_estimate_from_artifact,
        client_estimate_report_from_artifact,
        compose_client_estimate_markdown,
    )
    from core.hitl import get_draft_tz

    kg = KnowledgeRepository(db)
    draft = get_draft_tz(kg, project.id)
    estimate = client_estimate_from_artifact(draft)
    report = client_estimate_report_from_artifact(draft)
    if estimate is None:
        raise TzExportError("client estimate is not ready yet")
    markdown = compose_client_estimate_markdown(project, estimate, report)
    return export_markdown_file(markdown, f"{project.name}-smeta", fmt)


def export_tz_for_project(
    db: Session, project_id: uuid.UUID, fmt: TzExportFormat
) -> tuple[bytes, str, str]:
    from core.services import get_project

    project = get_project(db, project_id)
    if project is None:
        raise TzExportError("project not found")
    return export_tz_file(db, project, fmt)
